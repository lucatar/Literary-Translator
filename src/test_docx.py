from docx import Document

doc = Document()

p = doc.add_paragraph()
r = p.add_run("DEBUG")
r.italic = True

print("Before save:")
print(r._element.xml)

doc.save("debug.docx")

print("After save:")
print(r._element.xml)

doc.save("output/debug.docx")