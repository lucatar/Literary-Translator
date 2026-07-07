from docx import Document


class DocumentParser:

    def __init__(self, file_path: str):
        self.file_path = file_path

    def read(self) -> list[str]:
        """
        BACKWARD COMPATIBLE:
        pipeline továbbra is ezt használja
        """
        doc = Document(self.file_path)

        paragraphs = []

        for p in doc.paragraphs:
            text = p.text

            if text and text.strip():
                paragraphs.append(text.strip())

        return paragraphs

    # -------------------------------------------------
    # NEW
    # -------------------------------------------------

    def _extract_formatting(self, runs):

        formatting = []

        for r in runs:

            # csak akkor érdekes, ha VAN formázás
            if not (
                r.get("bold")
                or r.get("italic")
                or r.get("underline")
            ):
                continue

            formatting.append({
                "text": r["text"],
                "bold": r.get("bold", False),
                "italic": r.get("italic", False),
                "underline": r.get("underline", False)
            })

        return formatting

    def read_structured(self):
        """
        Structured paragraph representation.
        """

        doc = Document(self.file_path)

        structured = []

        for p in doc.paragraphs:

            runs = []
            full_text = ""

            for r in p.runs:

                if not r.text:
                    continue

                run = {
                    "text": r.text,
                    "bold": bool(r.bold),
                    "italic": bool(r.italic),
                    "underline": bool(r.underline),
                }

                runs.append(run)
                full_text += r.text

            if not full_text.strip():
                continue

            structured.append({
                "text": full_text.strip(),
                "runs": runs,                           # DOCX representation
                "formatting": self._extract_formatting(runs),  # internal representation
                "style": p.style.name if p.style else "Normal"
            })

        return structured