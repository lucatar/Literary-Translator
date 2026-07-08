from docx import Document


class DocxExporter:

    def __init__(self):
        print("DocxExporter __init__")

        self.doc = Document()

    def add_chunks(self, chunks):

        for chunk in chunks:

            if isinstance(chunk, str):
                self.doc.add_paragraph(chunk)
                continue

            for paragraph in chunk:

                text = paragraph.get("text", "")
                if not text:
                    continue

                p = self.doc.add_paragraph()

                try:
                    p.style = paragraph.get("style", "Normal")
                except Exception:
                    p.style = "Normal"

                formatting = paragraph.get("formatting", [])
                runs = paragraph.get("runs", [])

                #print("runs:", runs)
                print("formatting before:", formatting)
                # -----------------------
                # CASE 2: fallback runs
                # -----------------------
                if runs:

                    cursor = 0

                    for r in runs:

                        run_text = r.get("text", "")
                        if not run_text:
                            continue
                        #print(run_text)
                        run = p.add_run(run_text)

                        run.bold = r.get("bold")
                        run.italic = r.get("italic")
                        #print(run.italic)
                        #print("Assigned:", run.italic)
                        #print(run._element.xml)
                        run.underline = r.get("underline")

                    continue
                # -----------------------
                # CASE 1: formatting exists → APPLY AS OVERLAY
                # -----------------------
                """
                if formatting:
                    print("Beléptem a formatting-ba")
                    # 1. default run (full text)
                    formatted_run = p.add_run(text)

                    # apply paragraph-level formatting if present
                    f = formatting[0]
                    print("formatting after:", f)
                    formatted_run.bold = f.get("bold", False)
                    formatted_run.italic = f.get("italic", False)
                    formatted_run.underline = f.get("underline", False)

                    continue
                """
                # -----------------------
                # CASE 3: plain text
                # -----------------------
                p.add_run(text)

        self.doc.save("output/translated_style.docx")

    #def save(self, output_path: str):
        #self.doc.save(output_path)